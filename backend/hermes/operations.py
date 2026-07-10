"""工作区文件、联网、docx、draw.io、Figma 等原子操作。"""

from __future__ import annotations

import json
import re
import subprocess
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse

import requests

from backend.hermes.paths import get_workspace_root, resolve_workspace_path
from backend.hermes.settings import HermesSettings


def _domain_allowed(url: str, settings: HermesSettings) -> bool:
    if not settings.web_allowlist:
        return True
    host = (urlparse(url).hostname or "").lower()
    return any(host == d or host.endswith("." + d) for d in settings.web_allowlist)


def web_fetch_text(url: str, settings: HermesSettings) -> Dict[str, Any]:
    if not settings.web_fetch_enabled:
        return {"ok": False, "error": "HERMES_WEB_FETCH_ENABLED 未开启"}
    if not url.startswith(("http://", "https://")):
        return {"ok": False, "error": "仅支持 http(s) URL"}
    if not _domain_allowed(url, settings):
        return {"ok": False, "error": "域名不在 HERMES_WEB_ALLOWLIST 中"}
    try:
        r = requests.get(
            url,
            timeout=25,
            headers={"User-Agent": "emotional-chat-hermes/1.0"},
        )
        r.raise_for_status()
        body = r.content[: settings.web_max_bytes]
        ctype = (r.headers.get("content-type") or "").split(";")[0].strip().lower()
        if ctype in ("application/json",) or url.endswith(".json"):
            try:
                text = json.dumps(json.loads(body.decode("utf-8", errors="replace")), ensure_ascii=False)[:80000]
            except Exception:
                text = body.decode("utf-8", errors="replace")[:80000]
        else:
            text = body.decode("utf-8", errors="replace")[:80000]
        return {
            "ok": True,
            "url": url,
            "status": r.status_code,
            "content_type": ctype,
            "text_preview": text,
            "truncated": len(r.content) > len(body),
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


def list_workspace_files(
    workspace_root: str, relative_dir: str = "", glob_pattern: str = "*", limit: int = 200
) -> Dict[str, Any]:
    root = get_workspace_root(workspace_root)
    base = resolve_workspace_path(root, relative_dir or ".")
    if not base.is_dir():
        return {"ok": False, "error": "不是目录: {}".format(relative_dir)}
    out: List[str] = []
    for i, p in enumerate(sorted(base.rglob(glob_pattern))):
        if i >= limit:
            break
        if p.is_file():
            try:
                rel = str(p.relative_to(root))
            except ValueError:
                continue
            out.append(rel.replace("\\", "/"))
    return {"ok": True, "files": out, "count": len(out)}


def read_text_file(workspace_root: str, relative_path: str, max_chars: int = 120000) -> Dict[str, Any]:
    root = get_workspace_root(workspace_root)
    p = resolve_workspace_path(root, relative_path)
    if not p.is_file():
        return {"ok": False, "error": "文件不存在或不是文件"}
    if p.stat().st_size > max_chars + 1024:
        return {"ok": False, "error": "文件过大，请提高 max_chars 或换用分段读取"}
    text = p.read_text(encoding="utf-8", errors="replace")
    return {"ok": True, "path": str(p.relative_to(root)).replace("\\", "/"), "text": text[:max_chars], "truncated": len(text) > max_chars}


def write_text_file(workspace_root: str, relative_path: str, content: str) -> Dict[str, Any]:
    root = get_workspace_root(workspace_root)
    p = resolve_workspace_path(root, relative_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(content, encoding="utf-8")
    return {"ok": True, "path": str(p.relative_to(root)).replace("\\", "/"), "bytes": len(content.encode("utf-8"))}


def docx_get_text(workspace_root: str, relative_path: str) -> Dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        return {"ok": False, "error": "未安装 python-docx，请 pip install python-docx"}
    root = get_workspace_root(workspace_root)
    p = resolve_workspace_path(root, relative_path)
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error": "需要 .docx 文件"}
    doc = Document(str(p))
    paras = [para.text for para in doc.paragraphs]
    return {"ok": True, "paragraphs": paras, "joined": "\n".join(paras)[:200000]}


def docx_replace_text(
    workspace_root: str, relative_path: str, old_text: str, new_text: str
) -> Dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        return {"ok": False, "error": "未安装 python-docx"}
    root = get_workspace_root(workspace_root)
    p = resolve_workspace_path(root, relative_path)
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error": "需要 .docx 文件"}
    doc = Document(str(p))
    n = 0
    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)
                    n += 1
    doc.save(str(p))
    return {"ok": True, "replacements_applied": n, "path": str(p.relative_to(root)).replace("\\", "/")}


def docx_append_paragraph(workspace_root: str, relative_path: str, text: str) -> Dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        return {"ok": False, "error": "未安装 python-docx"}
    root = get_workspace_root(workspace_root)
    p = resolve_workspace_path(root, relative_path)
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error": "需要 .docx 文件"}
    doc = Document(str(p))
    doc.add_paragraph(text)
    doc.save(str(p))
    return {"ok": True, "path": str(p.relative_to(root)).replace("\\", "/"), "appended_chars": len(text)}


def drawio_set_cell_value(
    workspace_root: str, relative_path: str, cell_id: str, value: str
) -> Dict[str, Any]:
    """将 draw.io / mxGraph XML 中指定 id 的 mxCell 的 value 设为给定文本（会改写文件）。"""
    root = get_workspace_root(workspace_root)
    p = resolve_workspace_path(root, relative_path)
    if p.suffix.lower() not in (".drawio", ".xml"):
        return {"ok": False, "error": "需要 .drawio 或 XML"}
    tree = ET.parse(str(p))
    el_root = tree.getroot()
    hit = False
    for el in el_root.iter():
        if el.tag.endswith("mxCell") and el.attrib.get("id") == cell_id:
            el.set("value", value)
            hit = True
            break
    if not hit:
        return {"ok": False, "error": "未找到 id={} 的 mxCell".format(cell_id)}
    tree.write(str(p), encoding="utf-8", xml_declaration=True)
    return {"ok": True, "path": str(p.relative_to(root)).replace("\\", "/"), "cell_id": cell_id}


def drawio_summary(workspace_root: str, relative_path: str) -> Dict[str, Any]:
    root = get_workspace_root(workspace_root)
    p = resolve_workspace_path(root, relative_path)
    if p.suffix.lower() not in (".drawio", ".xml"):
        return {"ok": False, "error": "建议使用 .drawio 或含 mxfile 的 XML"}
    data = p.read_bytes()
    try:
        root_el = ET.fromstring(data)
    except ET.ParseError as e:
        return {"ok": False, "error": "XML 解析失败: {}".format(e)}
    cells: List[Dict[str, Any]] = []
    for el in root_el.iter():
        if el.tag.endswith("mxCell"):
            cid = el.attrib.get("id")
            val = (el.attrib.get("value") or "")[:200]
            if cid or val:
                cells.append({"id": cid, "value_preview": val})
            if len(cells) >= 400:
                break
    return {"ok": True, "cells": cells, "count": len(cells)}


def figma_file_summary(file_key: str, token: Optional[str]) -> Dict[str, Any]:
    if not token:
        return {"ok": False, "error": "未设置 FIGMA_ACCESS_TOKEN"}
    url = "https://api.figma.com/v1/files/{}".format(file_key.strip())
    try:
        r = requests.get(url, headers={"X-Figma-Token": token}, timeout=30)
        r.raise_for_status()
        data = r.json()
        name = data.get("name", "")
        comps = data.get("components", {}) or {}
        pages = data.get("document", {}).get("children", []) or []
        page_names = [c.get("name") for c in pages[:30]]
        return {
            "ok": True,
            "name": name,
            "page_names": page_names,
            "components_count": len(comps),
            "last_modified": data.get("lastModified"),
            "note": "只读元数据；画布级编辑请使用 Figma 客户端或官方 Plugin API。",
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


_SHELL_DENY = re.compile(
    r"\.\.|&&|\|\||;|`|\n|\r|powershell|pwsh|bash|sh\.exe|wsl|curl|wget|"
    r"certutil|bitsadmin|format\s|diskpart|\breg\s|mshta|rundll32|"
    r"invoke-expression|iex\b|downloadstring|net\s+user|schtasks|vbscript",
    re.I,
)


def run_shell_in_workspace(
    workspace_root: str, command: str, settings: HermesSettings
) -> Dict[str, Any]:
    """
    在 HERMES_WORKSPACE_ROOT 下以 cwd 执行 cmd 单行命令（需 HERMES_SHELL_ENABLED）。
    仍禁止管道、多语句与常见高危调用。
    """
    if not settings.shell_enabled:
        return {"ok": False, "error": "HERMES_SHELL_ENABLED 未开启"}
    cmd = (command or "").strip()
    if not cmd or len(cmd) > 4000:
        return {"ok": False, "error": "命令为空或过长"}
    if _SHELL_DENY.search(cmd):
        return {"ok": False, "error": "命令包含被禁止的模式（管道、多语句、powershell 等）"}
    root = get_workspace_root(workspace_root)
    try:
        proc = subprocess.run(
            ["cmd.exe", "/c", cmd],
            cwd=str(root),
            capture_output=True,
            text=True,
            timeout=settings.shell_timeout_sec,
            shell=False,
        )
        out = (proc.stdout or "") + (proc.stderr or "")
        return {
            "ok": proc.returncode == 0,
            "returncode": proc.returncode,
            "output": out[:24000],
            "truncated": len(out) > 24000,
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "命令超时"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


_URL_RE = re.compile(r"https?://[^\s\"'<>]+", re.I)
_PATH_RE = re.compile(
    r"(?:^|[\s\"'（(])([A-Za-z]:\\[^:*?\"<>|\r\n]+|(?:\.{0,2}[\\/])?[\w\-./\\]+\.(?:docx|drawio|xml))\b"
)


def extract_urls(text: str) -> List[str]:
    return _URL_RE.findall(text or "")


def extract_paths(text: str) -> List[str]:
    found = []
    for m in _PATH_RE.finditer(text or ""):
        g = m.group(1).strip()
        if g and g not in found:
            found.append(g)
    return found[:5]
