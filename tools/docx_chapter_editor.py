#!/usr/bin/env python3
"""Inspect and edit the first three chapters of the CodeX manuscript."""

from __future__ import annotations

import argparse
from pathlib import Path
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def iter_paragraphs(doc: Document):
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.replace("\n", " ").strip()
        if text:
            yield index, paragraph, text


def inspect(path: Path, start: int, end: int) -> None:
    doc = Document(path)
    for index, _paragraph, text in iter_paragraphs(doc):
        if start <= index <= end:
            print(f"{index}: {text}")


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_paragraph.style = style_name
    set_text(new_paragraph, text)
    return new_paragraph


def require_text(doc: Document, text: str) -> Paragraph:
    matches = [paragraph for _index, paragraph, value in iter_paragraphs(doc) if value == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph for {text!r}, found {len(matches)}")
    return matches[0]


def replace(doc: Document, old: str, new: str) -> None:
    set_text(require_text(doc, old), new)


def remove(doc: Document, text: str) -> None:
    delete_paragraph(require_text(doc, text))


def patch_cached_toc(path: Path) -> None:
    replacements = {
        "3.1　安装与与登录：先把入口跑通": "3.1　安装与登录：先把入口跑通",
        "3.1.1　安装Codex CLI": "3.1.1　安装 Codex CLI",
        "3.1.2　登录方式：ChatGPT账号或者API Key": "3.1.2　登录方式：ChatGPT 账号或者 API Key",
        "3.2.2　Web/Cloud：把任务交给云端环境": "3.2.3　Web/Cloud：把任务交给云端环境",
        "3.2.1 如何选择入口": "3.2.4　如何选择入口",
    }
    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    temp_path = path.with_suffix(".toc-patched.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                root = ET.fromstring(data)
                for paragraph in root.iter(w + "p"):
                    nodes = list(paragraph.iter(w + "t"))
                    joined = "".join(node.text or "" for node in nodes)
                    for old, new in replacements.items():
                        if joined.startswith(old) and nodes:
                            nodes[0].text = joined.replace(old, new, 1)
                            for node in nodes[1:]:
                                node.text = ""
                            break
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(info, data)
    temp_path.replace(path)


def edit(path: Path, output: Path) -> None:
    doc = Document(path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() == "AI 自主诊断并修复":
                        set_text(paragraph, "AI 协助诊断并尝试修复")
                    if paragraph.text.strip() == "代码变更 + 验证通过的 commit":
                        set_text(paragraph, "代码变更 + 验证证据 + 待审查说明")

    toolchain_intro = require_text(
        doc,
        "交付工具必须和工程工具链发生关系。Codex 的关键能力不只是生成代码，而是能够围绕代码调用工具，例如：",
    )
    insert_after(
        toolchain_intro,
        "代码清单 1-2：Codex 协同工程工具链的典型命令",
    )

    replace(
        doc,
        "把 Codex 看成交付工具，并不意味着要把它神化成一个“永远正确”的工程师。恰恰相反，只有承认它仍然会犯错，才能把它放进合适的工程位置。它能读仓库、能改文件、能跑命令，但这些能力解决的是“如何进入执行流程”的问题，并没有自动解决“它是否真正理解了业务意图”的问题。",
        "把 Codex 看成交付工具，并不意味着把它神化成“永远正确”的工程师。读仓库、改文件、跑命令，解决的是它如何进入执行流程；业务意图是否理解到位，仍要靠目标、上下文和审查来约束。",
    )
    replace(
        doc,
        "Codex 的错误往往不表现为语法层面的低级失败，而更容易出现在工程判断层面。例如，它可能根据现有代码推断出一个看似合理的实现，却没有意识到某个字段还要兼容历史数据；它可能顺利补齐单元测试，却没有覆盖跨服务调用中的真实异常；它也可能为了让测试通过而修改了过宽的范围，结果把一个局部任务扩展成了不必要的重构。对熟悉项目的人类工程师来说，这些问题未必难以发现，但如果完全把执行结果当成“正确答案”，风险就会被隐藏起来。",
        "它的偏差往往出现在工程判断层面：目标说得太宽，它会自行补全方向；项目里有未写出的兼容约束，它可能沿着局部代码做出不合适的实现；验证只覆盖了局部路径，它也可能把“已跑过检查”误当成“风险已经排尽”。这些问题未必立刻报错，却会抬高接受变更的成本。",
    )
    replace(
        doc,
        "从实践角度看，Codex 常见的偏差大致可以分为三类。",
        "因此，人类的注意力不该停在“它会不会写代码”，而要集中到四个判断上：任务是否定义清楚，项目上下文是否足够，验证证据是否可信，最终变更是否值得接受。",
    )
    replace(
        doc,
        "第一类是目标偏差。用户给出的任务本身过于宽泛，例如“优化这个模块”“把登录流程改好一点”，Codex 只能根据代码表象自行补全目标。它也许会做出技术上成立的改动，却不一定命中真正的产品诉求。",
        "这正是“方向不同”的含义。Codex 的进步，不在于把每次回答都变成标准答案，而在于把 AI 的作用位置从项目外部的建议者推进到工程流程内部的执行者。接下来要讨论的心智模型转变，正是如何在这个位置上与它协作。",
    )
    for text in [
        "第二类是上下文偏差。项目中许多约束并没有写进代码：某个接口虽然看似可以修改，实际上还有外部客户依赖；某个旧逻辑虽然不优雅，却是为了兼容历史迁移；某个测试没有覆盖的分支，恰恰是线上最敏感的路径。当这些信息没有进入文档、规则或任务描述时，Codex 很难凭空知道。",
        "第三类是验证偏差。测试通过说明某些行为得到了验证，却不等于所有风险都被排除。测试可能覆盖不足，命令可能只跑了局部检查，运行环境也可能和生产条件不同。Codex 能把验证前移，但它不能替代工程师对验证充分性的判断。",
        "这些局限并不否定 Codex 的价值，反而说明了它真正改变的是什么。过去使用对话式助手时，开发者常把大量时间花在搬运代码、修复路径、补接依赖和重新解释上下文上；使用 Codex 时，这些机械集成工作可以被大幅前移到智能体执行过程中。人类不再需要亲自搬运每个代码片段，而是把注意力更多放在四件事上：任务是否定义清楚，项目上下文是否足够，验证证据是否可信，最终变更是否值得接受。",
        "因此，Codex 更像一名执行力很强、反馈速度很快的新型协作者。它可以在短时间内完成阅读、修改、试错和汇报，把许多原本需要人工来回切换的工作连成闭环；但它仍需要明确的目标、稳定的规则和有判断力的审查。你不能把它当成只会回答问题的聊天机器人，也不应把它当成可以无人负责的自动程序员。",
        "这正是“方向不同”的含义。Codex 的进步不在于把每一次回答都变成标准答案，而在于把 AI 的作用位置从“项目外部的建议者”推进到“工程流程内部的执行者”。一旦 AI 进入执行流程，开发者的工作重心也随之改变：不再只是向模型索要一段代码，而是要学会定义任务、搭建边界、观察反馈、审查交付物。接下来要讨论的心智模型转变，正是从这里开始。",
    ]:
        remove(doc, text)

    rag_heading = require_text(doc, "2.2 Codex 完成任务时依赖什么")
    rag_heading.style = "Heading 2"
    replace(
        doc,
        "对于具备代码检索能力的编码 agent，可以把它处理超大代码库的方式近似理解为 RAG：先按任务检索相关上下文，再基于这些上下文进行分析和生成。Codex 的实际工作流还可能结合文件搜索、代码读取、工具调用和多轮验证，因此不一定等同于单一的标准 RAG 流程。",
        "可以把这种查阅过程类比为检索增强：先沿着任务线索找到相关材料，再把真正影响判断的内容带回当前工作上下文。",
    )
    replace(
        doc,
        '这种机制可以类比为“翻书”：上下文窗口像当前摊在桌面上的工作材料，容量有限；检索和工具调用则像去代码库或外部资料中查阅相关内容，再把有用的信息带回当前工作过程。上下文窗口是它的"工作记忆"——容量有限，但随时可用。如图2-1所示，RAG 是它的"查阅资料"能力——可以在更大的范围内搜索信息，但每次只能取回一小部分放入工作记忆。',
        "但类比不等于等同。真实编码任务还会结合文件搜索、代码读取、命令反馈和多轮验证；对读者更重要的判断是：检索能扩展视野，不能替代项目边界和验证证据。",
    )
    for text in [
        "检索过程通常基于三个信号：",
        '关键词匹配：根据任务描述中的关键词（如"用户"、"认证"、"JWT"）在代码库中搜索包含相同关键词的文件和代码片段',
        "语义相似度：使用向量嵌入（embedding）技术，将代码片段和任务描述都转换为数学向量，然后计算向量间的距离来衡量语义相关性",
    ]:
        remove(doc, text)

    replacements = {
        "3.1　安装与与登录：先把入口跑通": "3.1　安装与登录：先把入口跑通",
        "3.1.1　安装Codex CLI": "3.1.1　安装 Codex CLI",
        "3.1.2　登录方式：ChatGPT账号或者API Key": "3.1.2　登录方式：ChatGPT 账号或者 API Key",
        "3.2.2　Web/Cloud：把任务交给云端环境": "3.2.3　Web/Cloud：把任务交给云端环境",
        "3.2.1 如何选择入口": "3.2.4　如何选择入口",
        "读取你本地项目中的所有文件（不需要你手动提供）": "在授权范围内读取本地项目文件，减少手动粘贴上下文",
        "执行 git 操作（commit、branch、diff、log）": "配合 Git 查看差异、分支和历史，并按任务边界整理结果",
        "这条指令执行后，Codex 会自动：读取 README.md 和 pyproject.toml → 理解当前项目的依赖 → 在 README 中添加安装说明。整个过程你不需要做任何手动操作。": "这条指令执行后，Codex 可以沿着 README.md 和 pyproject.toml 获取上下文，在项目里形成一版安装说明变更。你的工作重点随之从手动搬运代码，转为审查范围、措辞和验证结果。",
        "第五步：执行测试，并处理真实环境反馈": "第四步：验证结果，并处理真实环境反馈",
        "图 3-4 codex CLI执行观察": "图 3-5 codex CLI执行观察",
        "图 3-5 codex CLI解决测试环境问题": "图 3-6 codex CLI解决测试环境问题",
        "图 3-6 codex CLI收敛测试结果": "图 3-7 codex CLI收敛测试结果",
        "图 3-7 codex CLI测试结果": "图 3-8 codex CLI测试结果",
        "图 3-8 codex CLI计划跑偏进行纠正": "图 3-9 codex CLI计划跑偏进行纠正",
        "图 3-9 codex CLI测试失败纠正失败点": "图 3-10 codex CLI测试失败纠正失败点",
        "如果 Codex 准备修改不相关文件，例如它想去改所有时间处理逻辑，而你只想改格式化函数，就应立即纠正，见图3-8。": "如果 Codex 准备修改不相关文件，例如它想去改所有时间处理逻辑，而你只想改格式化函数，就应立即纠正，见图3-9。",
        "如果测试失败，而 Codex 没抓住根因，就直接指出失败点，见图3-9。": "如果测试失败，而 Codex 没抓住根因，就直接指出失败点，见图3-10。",
    }
    for old, new in replacements.items():
        replace(doc, old, new)

    replace(
        doc,
        "在真实项目中，验证并不总会顺利抵达“测试通过”这一步。本次任务要求运行指定的单元测试。Codex 按要求发起验证后，首先遇到的并不是断言失败，如图3-5所示，而是测试环境本身尚未就绪：当前 Shell 无法找到 `pytest`，而项目配置又表明当前解释器版本低于项目要求。换用可用的 Python 环境后，测试继续向前推进，却在收集阶段被已有的模块导入问题阻塞。",
        "在真实项目中，验证并不总会顺利抵达“测试通过”这一步。本次任务要求运行指定的单元测试。Codex 按要求发起验证后，首先遇到的并不是断言失败，如图3-6所示，而是测试环境本身尚未就绪：当前 Shell 无法找到 `pytest`，而项目配置又表明当前解释器版本低于项目要求。换用可用的 Python 环境后，测试继续向前推进，却在收集阶段被已有的模块导入问题阻塞。",
    )
    replace(
        doc,
        "这类结果需要正确解读。图3-6说明 Codex 已经尝试执行验收步骤，但当前项目并不具备让这组测试完整跑通的条件。阻塞点既可能来自本地环境，也可能来自仓库中原本就存在的代码问题；它们与本次变更是否正确，并不是同一个判断。",
        "这类结果需要正确解读。图3-7说明 Codex 已经尝试执行验收步骤，但当前项目并不具备让这组测试完整跑通的条件。阻塞点既可能来自本地环境，也可能来自仓库中原本就存在的代码问题；它们与本次变更是否正确，并不是同一个判断。",
    )
    replace(
        doc,
        '如图3-7所示，这次修改严格控制在最小范围内：在 backend/core/utils/formatters.py 的 format_timestamp 中新增了format_type="date" 分支，使其返回 YYYY-MM-DD 格式日期字符串，并在 backend/tests/unit/test_core.py 现有的时间格式测试中补充了对应断言，原有 iso、rfc、unix、readable 以及自定义 strftime 回退逻辑都未改动。验证时先按要求运行 pytest backend/tests/unit/test_core.py -v，但当前环境没有可用的 pytest；随后切换到带有pytest 的 Python 3.12.12 环境继续执行，测试仍在收集阶段中断，先后暴露出 backend 包导入不可见以及现有代码中 backend.core.utils.interfaces 导入路径不存在的问题，因此新增断言尚未真正执行到；此外已运行 git diff--check，未发现补丁格式问题。',
        '如图3-8所示，这次修改严格控制在最小范围内：在 backend/core/utils/formatters.py 的 format_timestamp 中新增了 format_type="date" 分支，使其返回 YYYY-MM-DD 格式日期字符串，并在 backend/tests/unit/test_core.py 现有的时间格式测试中补充了对应断言，原有 iso、rfc、unix、readable 以及自定义 strftime 回退逻辑都未改动。验证时先按要求运行 pytest backend/tests/unit/test_core.py -v，但当前环境没有可用的 pytest；随后切换到带有 pytest 的 Python 3.12.12 环境继续执行，测试仍在收集阶段中断，先后暴露出 backend 包导入不可见以及现有代码中 backend.core.utils.interfaces 导入路径不存在的问题，因此新增断言尚未真正执行到；此外已运行 git diff --check，未发现补丁格式问题。',
    )
    inspire = require_text(doc, "启发")
    set_text(inspire, "第五步：整理结果，决定是否提交")
    inspire.style = "Heading 4"
    after = insert_after(
        inspire,
        "到了这一步，不应因为“改动很小”就直接提交。先把差异、已跑验证、未跑通原因和剩余风险写清楚；当结果可以被审查和接手时，再决定是否进入提交或后续修复。",
    )
    insight = insert_after(after, "启发", "Heading 4")
    insight.style = "Heading 4"

    replace(
        doc,
        "Codex 支持多个模型版本，不同模型在能力、速度和成本之间有不同的权衡。作为使用者，你需要根据任务类型选择合适的模型——不需要每次都纠结，但要有基本的判断框架。",
        "Codex 可用的模型、价格和产品默认值都会随版本演进而调整。正文里更值得保留的，是一个不容易过期的判断框架：先按任务风险、复杂度和等待时间选择能力档位，再在实际使用前核对当前可用模型与当期计费说明。",
    )
    replace(
        doc,
        "如表 3-1所示，Codex 的模型大致可以分为轻量级、标准级和高级三个层级。",
        "为便于第一次上手，可以先把模型选择近似看成三个能力档位，如表 3-1 所示。",
    )
    replace(
        doc,
        "默认用标准级：覆盖 80% 的日常任务。单文件修改、添加测试、简单重构、bug 修复——这些任务标准级模型完全够用。",
        "日常任务先选平衡档：单文件修改、补测试、边界明确的 bug 修复，通常先看是否能稳定理解仓库和给出验证证据。",
    )
    replace(
        doc,
        "简单任务用轻量级：格式化代码、添加注释/文档、生成样板文件、简单的配置修改。这些任务逻辑简单，轻量级模型的产出质量和标准级几乎没有差别，但速度快很多。",
        "低复杂度任务可选更快档：格式化、说明补充、样板文件和简单配置修改，更在意响应速度与成本控制。",
    )
    replace(
        doc,
        "复杂任务用高级：跨多个模块的重构、需要深度理解业务逻辑的实现、复杂的算法实现、涉及架构决策的改动。这些任务值得多花时间和成本来获得更好的结果。",
        "跨模块、高风险或需要长链路推理的任务，再考虑更强档位；即便如此，架构和业务决策也仍要由人把关。",
    )
    replace(
        doc,
        "一个实用的成本参考（以当前定价为准，实际价格可能调整）：",
        "成本判断也不必先背价格表。先看四个更稳定的驱动因素：上下文有多长、输出有多大、验证要跑多少轮、任务会不会反复重做。",
    )
    for text in [
        "简单任务（单文件修改、添加测试、格式化）：几百到几千 token，成本极低（通常不到 1 分钱人民币）",
        "中等任务（多文件修改 + 测试、API 实现、bug 修复）：几千到几万 token（通常几分钱到几毛钱人民币）",
        "复杂任务（大规模重构、多模块改动、架构调整）：几万到几十万 token（通常几毛钱到几块钱人民币）",
        "对于专业开发来说，这个成本几乎可以忽略不计。一个中级工程师的时薪是几百元，而 Codex 在几分钟内能完成可能需要几十分钟甚至几小时的手工工作。即使一个复杂任务花了几块钱，相比节省的工时，ROI 也是非常明确的。",
    ]:
        remove(doc, text)
    replace(
        doc,
        "但也应该有一些成本意识。以下是一些控制成本的实用技巧：",
        "有了这个顺序，再看控制成本就会更具体：",
    )
    replace(
        doc,
        '使用轻量级模型处理简单任务：格式化、注释等简单任务用轻量级模型，能节省 50-80% 的成本。',
        "简单任务别过度升档：低风险、低复杂度工作先选够用的能力档位，把更强模型留给真正需要推理深度的任务。",
    )
    replace(
        doc,
        "模型的响应速度也会影响你的工作节奏。轻量级模型通常在几秒内开始输出，适合需要快速迭代的场景；标准级模型的启动分析可能需要 10-30 秒，然后持续输出——这段时间不是浪费，Codex 在分析项目结构、理解代码依赖、规划执行步骤；高级模型的复杂任务可能需要 1-3 分钟的深度分析，但产出的质量通常更高。",
        "模型响应速度也会影响你的工作节奏。短任务更怕等待，长任务更怕方向错和验证不足。选择模型时，把“多久拿到第一版”与“返工一次要付出什么”放在一起看，会比只盯单次速度更稳。",
    )
    replace(
        doc,
        "一个实用的建议：不要盯着进度条看。在 Codex 工作时，你可以切换到其他任务（查邮件、看文档、喝杯咖啡），等它完成后再回来审查结果。",
        "一旦任务变长，人的节奏也要调整：把验收标准和中间检查点先写清，再回来审查结果，而不是把注意力耗在等待过程上。",
    )

    replace(
        doc,
        "Codex 会读取你指定路径下的文件内容，并将这些内容作为上下文发送给 AI 模型。如果你让它操作包含数据库密码、API 密钥、用户个人信息、商业机密等敏感数据的文件，这些数据就会离开你的本地环境，进入 AI 模型的处理流程。",
        "涉及数据库密码、API 密钥、用户个人信息和商业机密时，先按组织的数据分类、工作区策略和最小权限原则处理。能用脱敏样例、测试夹具和代码结构说明的问题，不要把真实敏感数据带进任务。",
    )
    replace(
        doc,
        "虽然 OpenAI 有数据隐私保护措施（API 数据不会被用于训练模型），但从工程安全的原则出发，你不应该让任何外部服务接触到生产环境的敏感数据。原因有三：第一，最小权限原则——AI 模型只需要代码上下文就能工作，不需要也不应该接触到数据库密码或用户数据；第二，供应链风险——即使 OpenAI 本身是可信的，数据传输过程中仍然存在被截获的风险（虽然概率极低）；第三，合规要求——很多行业（金融、医疗、政府）有严格的数据隔离要求，将生产数据发送给外部服务可能违反合规策略。",
        "具体的数据使用、保留与治理边界，应以当前官方说明和所在组织配置为准。对工程实践而言，先把真实生产数据排除在默认协作路径之外，通常比事后追问“这次是否越界”更稳。",
    )
    replace(
        doc,
        "可完全信任（直接使用）：代码格式化、添加注释和文档；编写单元测试和集成测试；简单的 CRUD 实现（增删改查）；Bug 修复（原因明确的）；依赖版本升级；代码翻译（不同语言之间的等价转换）。",
        "可快速验收：代码格式化、局部文档修订、边界明确的测试补充、样板代码和原因清楚的小修复。重点看范围是否收敛、验证是否匹配。",
    )
    replace(
        doc,
        "这个分级不是固定的——随着你的经验积累和对 Codex 的信任度提升，你可以适当放宽某些类别的信任范围。但在初期，保守一点没有坏处。",
        "这个分级不是固定的。随着项目规则、验证入口和团队经验逐渐成熟，验收路径可以更顺；但在第一次上手时，保守一些更容易建立正确习惯。",
    )
    replace(
        doc,
        "当你发现某个三步法模式反复出现时，就可以考虑将其自动化。Codex 提供了三个层次的自动化机制：",
        "当某个三步法模式反复出现时，再考虑把一部分步骤沉淀为更稳定的能力，例如可复用模板、自动检查或专门角色。它们的共同前提是：任务模式已经稳定，验收路径已经清楚。",
    )
    for text in [
        'Skill（技能）：如果某个任务模式经常出现（比如"添加一个新的 CRUD 模块"），可以将其封装为 Skill。Skill 是可复用的任务模板，包含预定义的上下文加载策略、任务描述模板和验收步骤。使用时只需一句话触发。',
        '钩子（Hook）：如果你希望在 Codex 执行的特定节点自动触发某些操作（比如每次写入文件后自动运行 formatter，每次 commit 前自动运行 lint），可以配置 Hook。钩子是事件驱动的自动化机制。',
        "子智能体（Subagent）：如果某个任务类型需要特殊的处理流程（比如代码审查需要专门的审查策略、安全扫描需要专门的检测规则），可以配置专门的子智能体。子智能体是角色化的自动化机制。",
        "这些高级功能将在后续章节详细讲解。现在你需要记住的是：好的工作流是从实践中提炼出来的。不要一开始就追求复杂的自动化，先用好基本的三步法。等你用了几十次、积累了足够的经验，你自然会感受到哪些环节重复度高、哪些步骤可以自动化。到时候再引入 Skill 和 Hook，水到渠成。",
        "过早的自动化是一个常见陷阱。有些工程师在用了几次 Codex 后就急着写 Skill 和 Hook，试图把一切都自动化。但因为没有足够的实践经验，他们设计的自动化流程往往不够灵活——能处理特定的任务模式，但稍有变化就失效了。好的自动化来自于对大量实际操作的模式提炼，而不是对少数几次操作的简单封装。",
    ]:
        remove(doc, text)
    insert_after(
        require_text(doc, "当某个三步法模式反复出现时，再考虑把一部分步骤沉淀为更稳定的能力，例如可复用模板、自动检查或专门角色。它们的共同前提是：任务模式已经稳定，验收路径已经清楚。"),
        "本章先记住这个顺序：先跑通最小闭环，再沉淀上下文、规则和验证。后续章节再展开 Skill、Hook 与子智能体。",
    )
    replace(
        doc,
        "模型选择不必纠结：轻量级（快而简）、标准级（平衡）、高级（强而慢），80% 的任务用标准级。成本在专业开发的语境下几乎可以忽略。更重要的是任务描述的质量——一个好的指令 + 标准模型，远胜过一个差的指令 + 高级模型。",
        "模型选择先看任务复杂度、风险和等待成本，再核对当期可用模型与计费说明。更重要的仍是任务描述质量：边界清楚、验证明确，往往比盲目升档更能减少返工。",
    )
    replace(
        doc,
        "知道禁区比知道能力更重要：敏感数据、生产环境、安全关键代码、不可逆操作、深度业务决策——这些不应该直接交给 Codex。分级信任原则是你的安全护栏：常规代码信任，业务逻辑抽查，安全架构必审。初期保守一点，随着经验积累再逐步放宽。",
        "知道禁区比知道能力更重要：敏感数据、生产环境、安全关键代码、不可逆操作、深度业务决策——这些不应该直接交给 Codex。分级信任原则是你的安全护栏：低风险变更可快速验收，业务逻辑要审，安全与架构必须深看。",
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    patch_cached_toc(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--inspect", action="store_true")
    parser.add_argument("--edit", action="store_true")
    parser.add_argument("--output", type=Path)
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--end", type=int, default=9999)
    args = parser.parse_args()

    if args.inspect:
        inspect(args.docx, args.start, args.end)
    if args.edit:
        if not args.output:
            parser.error("--edit requires --output")
        edit(args.docx, args.output)


if __name__ == "__main__":
    main()
